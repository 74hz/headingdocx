import re
import zipfile
from typing import List, Optional, Tuple

from docx import Document

from .heading_utils import (
    get_outline_level,
    is_bold,
    is_bold_and_large,
    is_bold_and_numbered,
    is_large,
    match_heading_patterns,
)


def get_headings(doc_path: str) -> List[Tuple[str, Optional[int]]]:
    """
    获取文档中所有标题目录，返回 [(标题文本, 级别), ...]
    智能识别标题级别，综合样式名、outlineLvl、加粗、字号、编号、文本模式等特征
    """
    doc = Document(doc_path)
    headings = []
    for paragraph in doc.paragraphs:
        style = getattr(paragraph, "style", None)
        style_name = getattr(style, "name", "") if style else ""
        text = paragraph.text.strip()
        level = None

        # 1. 样式名/ID判断
        if style_name:
            m1 = re.match(r"^Heading\s*(\d+)$", style_name, re.I)
            m2 = re.match(r"^标题\s*(\d+)$", style_name)
            if m1:
                level = int(m1.group(1))
            elif m2:
                level = int(m2.group(1))
            elif style_name.isdigit():
                num = int(style_name)
                if 1 <= num <= 9:
                    outline_lvl = get_outline_level(paragraph)
                    if outline_lvl is not None:
                        level = outline_lvl + 1
                    elif is_bold_and_large(paragraph, min_size=28):
                        level = num

        # 2. outlineLvl 属性
        if level is None:
            outline_lvl = get_outline_level(paragraph)
            if outline_lvl is not None:
                level = outline_lvl + 1

        # 3. 格式特征
        if level is None and text and len(text) < 200:
            if is_bold_and_large(paragraph, min_size=44):
                level = 1
            elif is_bold_and_large(paragraph, min_size=36):
                level = 2
            elif is_bold_and_large(paragraph, min_size=32):
                level = 3
            elif is_bold_and_numbered(paragraph):
                level = 2
            elif match_heading_patterns(text) and is_bold(paragraph):
                if is_large(paragraph, min_size=32):
                    level = 1
                else:
                    level = 2

        if level is not None:
            headings.append((text, level))
    return headings


def rebuild_doc_by_headings(doc_path: str, heading_texts: List[str], output_path: str):
    """
    根据给定标题文本列表（heading_texts）重组文档，生成新文档。
    heading_texts: 标题文本列表，按新顺序排列（可增删）
    output_path: 新文档保存路径
    """
    doc = Document(doc_path)
    new_doc = Document()
    # 提取所有标题及其内容块
    content_blocks = {}
    current_block = []
    current_heading = None
    for paragraph in doc.paragraphs:
        style = getattr(paragraph, "style", None)
        is_heading = (
            style is not None
            and getattr(style, "name", None)
            and style.name.startswith("Heading")
        )
        if is_heading:
            if current_heading:
                content_blocks[current_heading] = current_block
            current_heading = paragraph.text
            current_block = [paragraph]
        else:
            if current_heading:
                current_block.append(paragraph)
    if current_heading:
        content_blocks[current_heading] = current_block
    # 按新顺序组合
    for title in heading_texts:
        if title in content_blocks:
            for para in content_blocks[title]:
                new_doc.element.body.append(para._p)
    new_doc.save(output_path)


def get_paragraph_xml(doc_path: str) -> List[str]:
    """
    获取文档所有段落的XML字符串列表
    """
    doc = Document(doc_path)
    xml_list = []
    for paragraph in doc.paragraphs:
        xml_list.append(paragraph._p.xml)
    return xml_list


def regex_replace_in_docx(doc_path: str, pattern: str, repl: str, output_path: str):
    """
    只对 docx 正文（word/document.xml）做正则替换，并保存为新文件
    """
    with zipfile.ZipFile(doc_path, "r") as zin:
        with zipfile.ZipFile(output_path, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    text = re.sub(pattern, repl, text)
                    data = text.encode("utf-8")
                zout.writestr(item, data)
