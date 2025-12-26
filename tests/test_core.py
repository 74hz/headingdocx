import os
import sys

# 添加项目根目录到 sys.path，确保 headingdocx 可导入
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import pytest
from docx import Document

from headingdocx.core import (
    get_headings,
    get_paragraph_xml,
    rebuild_doc_by_headings,
    regex_replace_in_xml,
)

# 项目根目录
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "."))
TEST_DOC_PATH = os.path.join(ROOT, "input.docx")
OUTPUT_REBUILD_PATH = os.path.join(ROOT, "output_rebuild.docx")
OUTPUT_XML_PATH = os.path.join(ROOT, "output_xml.docx")
SAMPLE_DOCX_PATH = os.path.join(ROOT, "sample_test.docx")

# 配置：是否保留所有测试生成的文件（包括输出文件和样本文档）
KEEP_FILE = True  # 设置为 False 则测试后自动删除所有相关文件


def create_sample_docx(path: str):
    doc = Document()
    doc.add_heading("一级标题", level=1)
    doc.add_paragraph("正文内容1")
    doc.add_heading("二级标题", level=2)
    doc.add_paragraph("正文内容2")
    doc.add_heading("一级标题2", level=1)
    doc.add_paragraph("正文内容3")
    doc.save(path)


@pytest.fixture(scope="module")
def sample_docx():
    create_sample_docx(SAMPLE_DOCX_PATH)
    yield SAMPLE_DOCX_PATH
    if not KEEP_FILE and os.path.exists(SAMPLE_DOCX_PATH):
        os.remove(SAMPLE_DOCX_PATH)


def test_get_headings(sample_docx):
    headings = get_headings(sample_docx)
    assert isinstance(headings, list)
    assert len(headings) == 3
    assert headings[0][0] == "一级标题"
    assert headings[0][1] == 1
    assert headings[1][0] == "二级标题"
    assert headings[1][1] == 2
    assert headings[2][0] == "一级标题2"
    assert headings[2][1] == 1


def test_rebuild_doc_by_headings(sample_docx):
    headings = get_headings(sample_docx)
    heading_texts = [h[0] for h in headings[::-1]]  # 反转顺序
    rebuild_doc_by_headings(sample_docx, heading_texts, OUTPUT_REBUILD_PATH)
    assert os.path.exists(OUTPUT_REBUILD_PATH)
    doc = Document(OUTPUT_REBUILD_PATH)
    found_headings = [
        p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
    ]
    assert found_headings == heading_texts
    if not KEEP_FILE and os.path.exists(OUTPUT_REBUILD_PATH):
        os.remove(OUTPUT_REBUILD_PATH)


def test_get_paragraph_xml(sample_docx):
    xml_list = get_paragraph_xml(sample_docx)
    assert isinstance(xml_list, list)
    assert len(xml_list) > 0
    # 额外：将所有段落XML写入一个docx文件，便于人工检查
    doc = Document()
    for xml in xml_list:
        doc.add_paragraph(xml)
    doc.save(OUTPUT_XML_PATH)
    assert os.path.exists(OUTPUT_XML_PATH)
    if not KEEP_FILE and os.path.exists(OUTPUT_XML_PATH):
        os.remove(OUTPUT_XML_PATH)


def test_regex_replace_in_xml():
    xml = "<w:p><w:t>hello world</w:t></w:p>"
    pattern = r"hello"
    repl = "hi"
    result = regex_replace_in_xml(xml, pattern, repl)
    # 写入输出文件
    output_path = os.path.join(ROOT, "output_regex_replace.xml")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(result)
    assert "hi world" in result
    assert result.startswith("<w:p>")
    assert result.endswith("</w:p>")
    assert os.path.exists(output_path)
    if not KEEP_FILE and os.path.exists(output_path):
        os.remove(output_path)
